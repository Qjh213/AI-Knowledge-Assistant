import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.core.config import settings
from app.core.exceptions import (
    DocumentParseError,
    NoExtractableTextError,
)


@dataclass(frozen=True, slots=True)
class ParsedSection:
    text: str
    page_number: int | None
    metadata: dict[str, Any]


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    sections: tuple[ParsedSection, ...]
    character_count: int


class DocumentParserService:
    def __init__(
        self,
        storage_path: Path | None = None,
    ) -> None:
        self.storage_path = (
            storage_path or settings.document_storage_path
        ).resolve()

    def parse(self, file_path: str) -> ParsedDocument:
        path = self._resolve_path(file_path)
        extension = path.suffix.lower()

        try:
            if extension in {".txt", ".md"}:
                sections = self._parse_text(path)
            elif extension == ".pdf":
                sections = self._parse_pdf(path)
            elif extension == ".docx":
                sections = self._parse_docx(path)
            else:
                raise DocumentParseError(
                    path.name,
                    f"unsupported extension '{extension}'",
                )

        except (DocumentParseError, NoExtractableTextError):
            raise

        except Exception as exc:
            raise DocumentParseError(
                path.name,
                str(exc),
            ) from exc

        if not sections:
            raise NoExtractableTextError(path.name)

        character_count = sum(
            len(section.text)
            for section in sections
        )

        return ParsedDocument(
            sections=tuple(sections),
            character_count=character_count,
        )

    def _resolve_path(self, file_path: str) -> Path:
        path = (self.storage_path / file_path).resolve()

        if (
            path == self.storage_path
            or self.storage_path not in path.parents
        ):
            raise DocumentParseError(
                file_path,
                "invalid storage path",
            )

        if not path.is_file():
            raise DocumentParseError(
                file_path,
                "file does not exist",
            )

        return path

    def _parse_text(
        self,
        path: Path,
    ) -> list[ParsedSection]:
        text: str | None = None

        for encoding in ("utf-8-sig", "gb18030"):
            try:
                text = path.read_text(encoding=encoding)
                break
            except UnicodeDecodeError:
                continue

        if text is None:
            raise DocumentParseError(
                path.name,
                "unsupported text encoding",
            )

        cleaned = self._clean_text(text)

        if not cleaned:
            raise NoExtractableTextError(path.name)

        return [
            ParsedSection(
                text=cleaned,
                page_number=None,
                metadata={"extension": path.suffix.lower()},
            )
        ]

    def _parse_pdf(
        self,
        path: Path,
    ) -> list[ParsedSection]:
        reader = PdfReader(str(path))

        if reader.is_encrypted:
            raise DocumentParseError(
                path.name,
                "encrypted PDF files are not supported",
            )

        sections: list[ParsedSection] = []

        for page_index, page in enumerate(
            reader.pages,
            start=1,
        ):
            text = self._clean_text(page.extract_text() or "")

            if text:
                sections.append(
                    ParsedSection(
                        text=text,
                        page_number=page_index,
                        metadata={
                            "extension": ".pdf",
                            "page": page_index,
                        },
                    )
                )

        if not sections:
            raise NoExtractableTextError(path.name)

        return sections

    def _parse_docx(
        self,
        path: Path,
    ) -> list[ParsedSection]:
        document = DocxDocument(str(path))
        blocks: list[str] = []

        for paragraph in document.paragraphs:
            text = self._clean_text(paragraph.text)

            if text:
                blocks.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [
                    self._clean_text(cell.text)
                    for cell in row.cells
                ]
                row_text = " | ".join(
                    cell for cell in cells if cell
                )

                if row_text:
                    blocks.append(row_text)

        text = "\n\n".join(blocks).strip()

        if not text:
            raise NoExtractableTextError(path.name)

        return [
            ParsedSection(
                text=text,
                page_number=None,
                metadata={"extension": ".docx"},
            )
        ]

    @staticmethod
    def _clean_text(text: str) -> str:
        normalized = (
            text.replace("\x00", "")
            .replace("\r\n", "\n")
            .replace("\r", "\n")
        )
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)

        return normalized.strip()