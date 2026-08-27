from pathlib import Path

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter
from reportlab.pdfgen import canvas

from app.core.exceptions import (
    DocumentParseError,
    NoExtractableTextError,
)
from app.services.document_parser import DocumentParserService


def test_parse_utf8_text(tmp_path: Path) -> None:
    path = tmp_path / "utf8.txt"
    path.write_text(
        "第一段内容\n\n\n\n第二段内容",
        encoding="utf-8",
    )

    parser = DocumentParserService(storage_path=tmp_path)
    parsed = parser.parse(path.name)

    assert len(parsed.sections) == 1
    assert parsed.sections[0].text == "第一段内容\n\n第二段内容"
    assert parsed.sections[0].page_number is None
    assert parsed.character_count == len(
        "第一段内容\n\n第二段内容"
    )


def test_parse_gb18030_text(tmp_path: Path) -> None:
    path = tmp_path / "chinese.txt"
    path.write_bytes("中文编码测试".encode("gb18030"))

    parser = DocumentParserService(storage_path=tmp_path)
    parsed = parser.parse(path.name)

    assert parsed.sections[0].text == "中文编码测试"


def test_parse_markdown(tmp_path: Path) -> None:
    path = tmp_path / "guide.md"
    path.write_text(
        "# 标题\n\nMarkdown 内容",
        encoding="utf-8",
    )

    parser = DocumentParserService(storage_path=tmp_path)
    parsed = parser.parse(path.name)

    assert "# 标题" in parsed.sections[0].text
    assert parsed.sections[0].metadata["extension"] == ".md"


def test_parse_pdf(tmp_path: Path) -> None:
    path = tmp_path / "document.pdf"

    pdf = canvas.Canvas(str(path))
    pdf.drawString(72, 750, "PDF parser test content")
    pdf.save()

    parser = DocumentParserService(storage_path=tmp_path)
    parsed = parser.parse(path.name)

    assert len(parsed.sections) == 1
    assert "PDF parser test content" in parsed.sections[0].text
    assert parsed.sections[0].page_number == 1
    assert parsed.sections[0].metadata["page"] == 1


def test_parse_docx_paragraphs_and_tables(
    tmp_path: Path,
) -> None:
    path = tmp_path / "document.docx"

    document = DocxDocument()
    document.add_heading("Knowledge Assistant", level=1)
    document.add_paragraph("DOCX paragraph content")

    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Column A"
    table.cell(0, 1).text = "Column B"

    document.save(path)

    parser = DocumentParserService(storage_path=tmp_path)
    parsed = parser.parse(path.name)

    text = parsed.sections[0].text

    assert "Knowledge Assistant" in text
    assert "DOCX paragraph content" in text
    assert "Column A | Column B" in text
    assert parsed.sections[0].page_number is None


def test_reject_empty_text_document(
    tmp_path: Path,
) -> None:
    path = tmp_path / "empty.txt"
    path.write_text("   \n\n", encoding="utf-8")

    parser = DocumentParserService(storage_path=tmp_path)

    with pytest.raises(NoExtractableTextError):
        parser.parse(path.name)


def test_reject_pdf_without_extractable_text(
    tmp_path: Path,
) -> None:
    path = tmp_path / "blank.pdf"

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)

    with path.open("wb") as output:
        writer.write(output)

    parser = DocumentParserService(storage_path=tmp_path)

    with pytest.raises(NoExtractableTextError):
        parser.parse(path.name)


def test_reject_parser_path_traversal(
    tmp_path: Path,
) -> None:
    storage_path = tmp_path / "storage"
    storage_path.mkdir()

    outside = tmp_path / "outside.txt"
    outside.write_text("outside", encoding="utf-8")

    parser = DocumentParserService(
        storage_path=storage_path,
    )

    with pytest.raises(DocumentParseError):
        parser.parse("../outside.txt")


def test_reject_missing_document(
    tmp_path: Path,
) -> None:
    parser = DocumentParserService(storage_path=tmp_path)

    with pytest.raises(DocumentParseError):
        parser.parse("missing.txt")